from docx import Document
import pandas as pd

# Load the CSV file
file_path = str(input("Enter the full path and filename for CSV: "))
df_quiz = pd.read_csv(file_path)

# Loading the document
doc = Document()
doc.add_heading('Quiz Questions and Answers', 0)

# Extract unique questions
questions = df_quiz['Q Text'].unique()

for question in questions:
    # Add question to the document
    doc.add_paragraph(question, style='Heading 2')

    # Extract possible answers for the question
    answers = df_quiz[df_quiz['Q Text'] == question]

    for _, row in answers.iterrows():
        answer = row['Answer']
        is_correct = row['Answer Match'] == 'Checked'

        # Add the answer into the document, BOLD if correct answer
        if is_correct:
            doc.add_paragraph(f" - {answer}").bold = True
        else:
            doc.add_paragraph(f" - {answer}")

# Save the document
output_path = str(input("Enter the full path with file name for .DOCX output: "))
doc.save(output_path)
